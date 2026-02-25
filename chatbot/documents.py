from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
import os
from pathlib import Path
import re

from docx import Document
from pypdf import PdfReader


WHITESPACE_PATTERN = re.compile(r"\s+")
SLUG_PATTERN = re.compile(r"[^a-zA-Z0-9]+")
SUPPORTED_EXTENSIONS = {".docx", ".pdf"}
DEFAULT_OCR_LANGUAGE = "eng"
DEFAULT_OCR_SCALE = 2.0
MIN_DIRECT_PDF_TEXT_CHARS = 20
DEFAULT_TESSERACT_PATHS = (
    Path("C:/Program Files/Tesseract-OCR/tesseract.exe"),
    Path("C:/Program Files (x86)/Tesseract-OCR/tesseract.exe"),
)


@dataclass(frozen=True)
class DocumentChunk:
    chunk_id: str
    source_name: str
    text: str


def _normalize_text(value: str) -> str:
    return WHITESPACE_PATTERN.sub(" ", value).strip()


def _slugify(value: str) -> str:
    slug = SLUG_PATTERN.sub("-", value.strip().lower()).strip("-")
    return slug or "document"


def _has_enough_text(value: str, min_chars: int = MIN_DIRECT_PDF_TEXT_CHARS) -> bool:
    meaningful_chars = sum(character.isalnum() for character in value)
    return meaningful_chars >= min_chars


def _resolve_tesseract_cmd() -> str:
    configured_path = os.getenv("TESSERACT_CMD", "").strip()
    if configured_path:
        return configured_path

    for default_path in DEFAULT_TESSERACT_PATHS:
        if default_path.exists():
            return str(default_path)

    return "tesseract"


def extract_text_from_docx(file_bytes: bytes) -> str:
    document = Document(BytesIO(file_bytes))
    blocks: list[str] = []

    for paragraph in document.paragraphs:
        cleaned = _normalize_text(paragraph.text)
        if cleaned:
            blocks.append(cleaned)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cleaned = _normalize_text(cell.text)
                if cleaned:
                    blocks.append(cleaned)

    return "\n".join(blocks)


def _extract_text_from_pdf_with_ocr(
    file_bytes: bytes,
    page_indices: list[int],
    language: str,
    scale: float,
) -> dict[int, str]:
    try:
        import pypdfium2 as pdfium
    except ImportError as exc:  # pragma: no cover - runtime dependency
        raise RuntimeError(
            "OCR support requires pypdfium2. Install dependencies with "
            "`pip install -r requirements.txt`."
        ) from exc

    try:
        import pytesseract
        from pytesseract import TesseractNotFoundError
    except ImportError as exc:  # pragma: no cover - runtime dependency
        raise RuntimeError(
            "OCR support requires pytesseract. Install dependencies with "
            "`pip install -r requirements.txt`."
        ) from exc

    pytesseract.pytesseract.tesseract_cmd = _resolve_tesseract_cmd()

    try:
        pytesseract.get_tesseract_version()
    except Exception as exc:  # pragma: no cover - machine specific
        raise RuntimeError(
            "Tesseract OCR engine is not installed or not on PATH. "
            "Install Tesseract and optionally set TESSERACT_CMD in .env."
        ) from exc

    ocr_text_by_page: dict[int, str] = {}
    document = pdfium.PdfDocument(file_bytes)

    try:
        for page_index in page_indices:
            page = document.get_page(page_index)
            try:
                bitmap = page.render(scale=scale)
                try:
                    page_image = bitmap.to_pil()
                finally:
                    bitmap.close()
            finally:
                page.close()

            try:
                page_text = pytesseract.image_to_string(page_image, lang=language)
            except TesseractNotFoundError as exc:  # pragma: no cover - machine specific
                raise RuntimeError(
                    "Tesseract OCR engine is not installed or not on PATH. "
                    "Install Tesseract and optionally set TESSERACT_CMD in .env."
                ) from exc

            cleaned = _normalize_text(page_text)
            if cleaned:
                ocr_text_by_page[page_index] = cleaned
    finally:
        document.close()

    return ocr_text_by_page


def extract_text_from_pdf(
    file_bytes: bytes,
    enable_ocr: bool = True,
    ocr_language: str | None = None,
    ocr_scale: float = DEFAULT_OCR_SCALE,
) -> str:
    reader = PdfReader(BytesIO(file_bytes))
    page_texts: list[str] = []
    pages_for_ocr: list[int] = []

    for page_index, page in enumerate(reader.pages):
        page_text = page.extract_text() or ""
        cleaned = _normalize_text(page_text)
        page_texts.append(cleaned)

        if enable_ocr and not _has_enough_text(cleaned):
            pages_for_ocr.append(page_index)

    if pages_for_ocr:
        language = (ocr_language or os.getenv("OCR_LANGUAGE", DEFAULT_OCR_LANGUAGE)).strip()
        language = language or DEFAULT_OCR_LANGUAGE

        try:
            ocr_text_by_page = _extract_text_from_pdf_with_ocr(
                file_bytes=file_bytes,
                page_indices=pages_for_ocr,
                language=language,
                scale=ocr_scale,
            )
        except RuntimeError:
            if not any(_has_enough_text(text, min_chars=1) for text in page_texts):
                raise
        else:
            for page_index, ocr_text in ocr_text_by_page.items():
                if _has_enough_text(ocr_text, min_chars=1):
                    page_texts[page_index] = ocr_text

    return "\n".join(text for text in page_texts if text)


def extract_text_from_file(
    file_name: str,
    file_bytes: bytes,
    enable_pdf_ocr: bool = True,
) -> str:
    extension = Path(file_name).suffix.lower()

    if extension == ".docx":
        return extract_text_from_docx(file_bytes)
    if extension == ".pdf":
        return extract_text_from_pdf(file_bytes, enable_ocr=enable_pdf_ocr)

    supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
    raise ValueError(f"Unsupported file type '{extension}'. Supported types: {supported}")


def chunk_text(text: str, chunk_size: int = 220, chunk_overlap: int = 40) -> list[str]:
    if chunk_size <= 0:
        raise ValueError("chunk_size must be greater than zero")
    if chunk_overlap < 0:
        raise ValueError("chunk_overlap cannot be negative")
    if chunk_overlap >= chunk_size:
        raise ValueError("chunk_overlap must be smaller than chunk_size")

    words = text.split()
    if not words:
        return []

    step = chunk_size - chunk_overlap
    chunks: list[str] = []

    for start in range(0, len(words), step):
        end = min(start + chunk_size, len(words))
        chunk_words = words[start:end]
        if not chunk_words:
            continue
        chunks.append(" ".join(chunk_words))
        if end == len(words):
            break

    return chunks


def _build_chunks_from_text(
    file_name: str,
    text: str,
    chunk_size: int = 220,
    chunk_overlap: int = 40,
) -> list[DocumentChunk]:
    text_chunks = chunk_text(text, chunk_size=chunk_size, chunk_overlap=chunk_overlap)

    source_slug = _slugify(file_name)
    chunks: list[DocumentChunk] = []

    for index, chunk_value in enumerate(text_chunks, start=1):
        chunks.append(
            DocumentChunk(
                chunk_id=f"{source_slug}-chunk-{index}",
                source_name=file_name,
                text=chunk_value,
            )
        )

    return chunks


def build_chunks_from_file(
    file_name: str,
    file_bytes: bytes,
    chunk_size: int = 220,
    chunk_overlap: int = 40,
    enable_pdf_ocr: bool = True,
) -> list[DocumentChunk]:
    text = extract_text_from_file(
        file_name=file_name,
        file_bytes=file_bytes,
        enable_pdf_ocr=enable_pdf_ocr,
    )
    return _build_chunks_from_text(
        file_name=file_name,
        text=text,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
    )


def build_chunks_from_docx(
    file_name: str,
    file_bytes: bytes,
    chunk_size: int = 220,
    chunk_overlap: int = 40,
) -> list[DocumentChunk]:
    text = extract_text_from_docx(file_bytes)
    return _build_chunks_from_text(
        file_name=file_name,
        text=text,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
    )
