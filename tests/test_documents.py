from __future__ import annotations

from io import BytesIO
from pathlib import Path

import pytest
from docx import Document as WordDocument

from chatbot import documents
from chatbot.documents import DocumentChunk


class _FakePage:
    def __init__(self, text: str) -> None:
        self._text = text

    def extract_text(self) -> str:
        return self._text


def _make_reader(page_texts: list[str]):
    class _FakeReader:
        def __init__(self, _stream: BytesIO) -> None:
            self.pages = [_FakePage(text) for text in page_texts]

    return _FakeReader


def test_normalize_text_collapses_whitespace() -> None:
    assert documents._normalize_text("  alpha\n\tbeta   gamma  ") == "alpha beta gamma"


def test_slugify_returns_document_when_empty() -> None:
    assert documents._slugify("***") == "document"


def test_has_enough_text_counts_alphanumeric_characters() -> None:
    assert documents._has_enough_text("A1 B2 C3", min_chars=6)
    assert not documents._has_enough_text("---", min_chars=1)


def test_resolve_tesseract_cmd_prefers_environment(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv("TESSERACT_CMD", r"C:\custom\tesseract.exe")
    assert documents._resolve_tesseract_cmd() == r"C:\custom\tesseract.exe"


def test_resolve_tesseract_cmd_uses_detected_default(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    fake_path = tmp_path / "tesseract.exe"
    fake_path.write_text("", encoding="utf-8")
    monkeypatch.delenv("TESSERACT_CMD", raising=False)
    monkeypatch.setattr(documents, "DEFAULT_TESSERACT_PATHS", (fake_path,))

    assert documents._resolve_tesseract_cmd() == str(fake_path)


def test_extract_text_from_docx_reads_paragraphs_and_table_cells() -> None:
    document = WordDocument()
    document.add_paragraph("Paragraph text")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Table text"

    buffer = BytesIO()
    document.save(buffer)

    text = documents.extract_text_from_docx(buffer.getvalue())

    assert "Paragraph text" in text
    assert "Table text" in text


def test_extract_text_from_file_dispatches_by_extension(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(documents, "extract_text_from_docx", lambda _bytes: "docx-text")
    monkeypatch.setattr(
        documents,
        "extract_text_from_pdf",
        lambda _bytes, enable_ocr=True: f"pdf-text-{enable_ocr}",
    )

    assert documents.extract_text_from_file("notes.docx", b"x") == "docx-text"
    assert (
        documents.extract_text_from_file("scan.PDF", b"x", enable_pdf_ocr=False)
        == "pdf-text-False"
    )


def test_extract_text_from_file_rejects_unsupported_extension() -> None:
    with pytest.raises(ValueError, match="Unsupported file type"):
        documents.extract_text_from_file("data.txt", b"hello")


@pytest.mark.parametrize(
    ("chunk_size", "chunk_overlap", "message"),
    [
        (0, 0, "chunk_size must be greater than zero"),
        (5, -1, "chunk_overlap cannot be negative"),
        (5, 5, "chunk_overlap must be smaller than chunk_size"),
    ],
)
def test_chunk_text_validates_inputs(
    chunk_size: int,
    chunk_overlap: int,
    message: str,
) -> None:
    with pytest.raises(ValueError, match=message):
        documents.chunk_text("one two three", chunk_size=chunk_size, chunk_overlap=chunk_overlap)


def test_chunk_text_creates_overlapping_chunks() -> None:
    text = "one two three four five six seven eight nine ten"
    chunks = documents.chunk_text(text, chunk_size=4, chunk_overlap=1)

    assert chunks == [
        "one two three four",
        "four five six seven",
        "seven eight nine ten",
    ]


def test_build_chunks_from_text_generates_stable_ids() -> None:
    chunks = documents._build_chunks_from_text(
        file_name="My File.docx",
        text="alpha beta gamma delta epsilon",
        chunk_size=3,
        chunk_overlap=1,
    )

    assert [chunk.chunk_id for chunk in chunks] == ["my-file-docx-chunk-1", "my-file-docx-chunk-2"]
    assert all(chunk.source_name == "My File.docx" for chunk in chunks)


def test_build_chunks_from_file_forwards_ocr_flag(monkeypatch: pytest.MonkeyPatch) -> None:
    captured: dict[str, object] = {}

    def fake_extract_text_from_file(
        file_name: str,
        file_bytes: bytes,
        enable_pdf_ocr: bool,
    ) -> str:
        captured["file_name"] = file_name
        captured["file_bytes"] = file_bytes
        captured["enable_pdf_ocr"] = enable_pdf_ocr
        return "one two three four"

    monkeypatch.setattr(documents, "extract_text_from_file", fake_extract_text_from_file)

    chunks = documents.build_chunks_from_file(
        file_name="sample.pdf",
        file_bytes=b"pdf",
        chunk_size=3,
        chunk_overlap=1,
        enable_pdf_ocr=False,
    )

    assert captured["enable_pdf_ocr"] is False
    assert isinstance(chunks[0], DocumentChunk)


def test_extract_text_from_pdf_skips_ocr_when_direct_text_is_sufficient(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(
        documents,
        "PdfReader",
        _make_reader([
            "This page contains enough extracted text to avoid OCR 1234567890.",
            "Second page also has direct content 1234567890.",
        ]),
    )

    def fail_ocr(**_kwargs: object) -> dict[int, str]:
        raise AssertionError("OCR should not be called for text-rich pages")

    monkeypatch.setattr(documents, "_extract_text_from_pdf_with_ocr", fail_ocr)

    extracted = documents.extract_text_from_pdf(b"pdf-bytes", enable_ocr=True)

    assert "avoid OCR" in extracted
    assert "Second page" in extracted


def test_extract_text_from_pdf_uses_ocr_for_low_text_pages(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(
        documents,
        "PdfReader",
        _make_reader([
            "",
            "this page has enough direct text to stay as-is 1234567890",
        ]),
    )
    monkeypatch.setenv("OCR_LANGUAGE", "eng")

    captured: dict[str, object] = {}

    def fake_ocr(
        file_bytes: bytes,
        page_indices: list[int],
        language: str,
        scale: float,
    ) -> dict[int, str]:
        captured["file_bytes"] = file_bytes
        captured["page_indices"] = page_indices
        captured["language"] = language
        captured["scale"] = scale
        return {0: "OCR recovered text"}

    monkeypatch.setattr(documents, "_extract_text_from_pdf_with_ocr", fake_ocr)

    extracted = documents.extract_text_from_pdf(b"pdf-bytes", enable_ocr=True, ocr_scale=3.0)

    assert captured["page_indices"] == [0]
    assert captured["language"] == "eng"
    assert captured["scale"] == 3.0
    assert "OCR recovered text" in extracted
    assert "stay as-is" in extracted


def test_extract_text_from_pdf_raises_when_ocr_fails_and_no_text_available(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(documents, "PdfReader", _make_reader(["", "   "]))
    monkeypatch.setattr(
        documents,
        "_extract_text_from_pdf_with_ocr",
        lambda **_kwargs: (_ for _ in ()).throw(RuntimeError("OCR unavailable")),
    )

    with pytest.raises(RuntimeError, match="OCR unavailable"):
        documents.extract_text_from_pdf(b"pdf-bytes", enable_ocr=True)


def test_extract_text_from_pdf_keeps_partial_direct_text_when_ocr_fails(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(
        documents,
        "PdfReader",
        _make_reader(["Readable direct page 12345", ""]),
    )
    monkeypatch.setattr(
        documents,
        "_extract_text_from_pdf_with_ocr",
        lambda **_kwargs: (_ for _ in ()).throw(RuntimeError("OCR unavailable")),
    )

    extracted = documents.extract_text_from_pdf(b"pdf-bytes", enable_ocr=True)

    assert extracted == "Readable direct page 12345"
